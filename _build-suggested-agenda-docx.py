"""
Rebuild suggested-agenda-therabreath.docx — short shareable agenda for TheraBreath.
Run: python _build-suggested-agenda-docx.py
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "booklet" / "documents" / "suggested-agenda-therabreath.docx"

ITEMS = [
    ("9:00 AM", "Arrive — conference room across from The Flavor Factory"),
    (
        "9:00–10:00",
        "Presentation — four pillars (resiliency, innovation, operations, partnership); "
        "Lakewood site parity; competitive & functional innovation",
    ),
    ("10:00–10:15", "Facility walkthrough — production, QC, TheraBreath room"),
    ("10:15–10:30", "Break"),
    (
        "10:30–11:15",
        "Tasting — ten workshop prototypes; blind flavor mapping & prototype scorecard",
    ),
    ("11:30–1:00", "Lunch off site — hosts take the group"),
    (
        "1:00 PM+",
        "Open Q&A — procurement, growth, chlorite flavor work, Lakewood validation",
    ),
]


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Suggested Agenda")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "Breath of Innovation · The Flavor Factory × TheraBreath\n"
        "Tuesday, July 8, 2026 · Norco, California"
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x5C, 0x66, 0x78)

    note = doc.add_paragraph()
    note_run = note.add_run("All times below are suggested.")
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x5C, 0x66, 0x78)

    doc.add_paragraph()

    for time, subject in ITEMS:
        p = doc.add_paragraph(style="List Bullet")
        time_run = p.add_run(f"{time} — ")
        time_run.bold = True
        time_run.font.color.rgb = RGBColor(0x00, 0x8F, 0xD3)
        p.add_run(subject)

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Questions before July 8? Contact The Flavor Factory team.\n"
        "(951) 273-9877 · flavorfactory.net\n"
        "2058 Second Street, Norco, CA 92860"
    )
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x5C, 0x66, 0x78)

    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    build()