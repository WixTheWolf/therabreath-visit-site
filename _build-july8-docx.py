"""Rebuild july-8-briefing.docx — editable briefing matching the PDF."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from _july8_briefing_content import (
    AGENDA_ROWS,
    ARRIVE_NOTE,
    EMAIL_SUBJECT,
    INNOVATION_ITEMS,
    OPENING,
    OUTCOMES_TB,
    PILLARS,
    STRATEGIC_QUESTIONS,
    TOPICS_COMPACT,
)

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "booklet" / "documents" / "july-8-briefing.docx"
TFF_LOGO = ROOT / "assets" / "companies" / "logos" / "tff-logo.png"
TB_LOGO = ROOT / "assets" / "companies" / "logos" / "therabreath-logo.png"

BLUE = RGBColor(0x00, 0x8F, 0xD3)
INK = RGBColor(0x0A, 0x16, 0x28)
MUTED = RGBColor(0x5C, 0x66, 0x78)
GREEN = RGBColor(0x5F, 0xB8, 0x32)


def _style_normal(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def _add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = BLUE


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = INK


def build() -> None:
    doc = Document()
    _style_normal(doc)

    header = doc.add_table(rows=1, cols=3)
    header.autofit = True
    left, mid, right = header.rows[0].cells
    if TFF_LOGO.is_file():
        left.paragraphs[0].add_run().add_picture(str(TFF_LOGO), height=Inches(0.32))
    mid.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid_run = mid.paragraphs[0].add_run("×")
    mid_run.bold = True
    mid_run.font.size = Pt(14)
    mid_run.font.color.rgb = BLUE
    if TB_LOGO.is_file():
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(TB_LOGO), height=Inches(0.32))

    doc.add_paragraph()

    _add_kicker(doc, "BREATH OF INNOVATION")
    _add_heading(doc, "Your July 8 briefing")
    sub = doc.add_paragraph("Tuesday, July 8, 2026 · Norco, California")
    sub.runs[0].font.color.rgb = MUTED
    doc.add_paragraph(OPENING)

    arrive = doc.add_paragraph()
    arrive_run = arrive.add_run(
        ARRIVE_NOTE.replace("<b>", "").replace("</b>", "")
    )
    arrive_run.bold = True

    _add_kicker(doc, "YOUR DAY")
    table = doc.add_table(rows=1 + len(AGENDA_ROWS), cols=3)
    table.style = "Table Grid"
    headers = ("Time", "Session", "Focus")
    for col, label in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = INK
    for row_idx, (time, session, focus) in enumerate(AGENDA_ROWS, start=1):
        table.rows[row_idx].cells[0].text = time
        table.rows[row_idx].cells[1].text = session
        table.rows[row_idx].cells[2].text = focus
        for run in table.rows[row_idx].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLUE
        for run in table.rows[row_idx].cells[1].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    _add_kicker(doc, "FOUR PILLARS · 9:30")
    for title, desc in PILLARS:
        p = doc.add_paragraph(style="List Bullet")
        t = p.add_run(f"{title} — ")
        t.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    _add_kicker(doc, "TOPICS WE WILL COVER")
    for title, desc in TOPICS_COMPACT:
        p = doc.add_paragraph(style="List Bullet")
        t = p.add_run(f"{title} — ")
        t.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    _add_kicker(doc, "INNOVATION · FLAVOR PARTNER")
    for item in INNOVATION_ITEMS:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()
    _add_kicker(doc, "TEN QUESTIONS FOR YOUR TEAM")
    for i, question in enumerate(STRATEGIC_QUESTIONS, start=1):
        p = doc.add_paragraph(style="List Number")
        num = p.add_run(f"{i}. ")
        num.bold = True
        num.font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)
        p.add_run(question)

    doc.add_paragraph()
    _add_kicker(doc, "WHAT YOU SHOULD LEAVE WITH")
    for item in OUTCOMES_TB:
        doc.add_paragraph(f"✓  {item}", style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "The Flavor Factory · 2058 Second Street, Norco, CA 92860 · (951) 273-9877\n"
        f"{EMAIL_SUBJECT}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
